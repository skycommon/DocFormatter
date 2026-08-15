# -*- coding: utf-8 -*-
r"""
文档快速排版 —— 核心排版引擎（纯离线启发式，不依赖大模型）
负责：规范化文本 -> 识别结构（标题/段落/列表/表格）-> 整理重点加粗 -> 生成 Word(.docx)

设计原则：
- 所有解析都是纯函数，便于离线单测。
- Markdown 标题 (`#`~`######`) 永远被尊重。
- Markdown 表格（GFM）自动识别：表头 + 分隔行（`:---`/`:--:`/`---:` 控制对齐）。
- 普通文本的「智能标题识别」为可选开关，且采取保守策略，避免误判。
- 重点加粗只处理安全模式：Markdown `**x**` 与「术语：内容」中冒号前的短术语。
- 数学公式（LaTeX）：块级 `$$...$$` / `\[...\]` 与行内 `$...$` / `\(...\)` 全程保护，
  不被加粗 / 标题 / 段落合并规则破坏；「公式美化」开关控制是否去括号并斜体显示。
- 正文引用标注（顺序编码制 `[1]`）：默认渲染为上标，仅改变字形、不改动数字与方括号
  内容；公式片段已被隔离，不会误伤；可在 GUI 关闭。
"""

import re
import datetime

# ---------------------------------------------------------------------------
# 常量
# ---------------------------------------------------------------------------

# 带圈数字 ①~⑳（参考文献「带圈数字」样式用；超过 20 自动回退方括号）
CIRCLED = "①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳"

# 中文句末标点（用于判断一行是否像「完整句子」而非标题）
CN_SENT_END = "。！？”」』）"
# 中文/英文冒号（用于「术语：内容」加粗）
COLON_RE = re.compile(r"[：:]")
# 普通文本中应避免当作标题的结尾标点
WEAK_END = "。，；：！？、…"

# 列表项标记
BULLET_RE = re.compile(r"^\s*([\-\*\u2022\u00b7\u2014\u25aa\u25cf\u25cb\u25a0\+])\s+(.*)$")
# 有序列表标记：1. 1、 (1) ① 等
ORDERED_RE = re.compile(r"^\s*([0-9]+[\.、]|\([0-9]+\)|[\(（][0-9]+[\)）]|[①②③④⑤⑥⑦⑧⑨⑩])\s*(.*)$")
# Markdown 标题
MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
# 章节 / 编号标题
CHAPTER_RE = re.compile(r"^第[一二三四五六七八九十百千0-9]+章")
DOTNUM_RE = re.compile(r"^[0-9]+(\.[0-9]+)*[\.、]\s*")
CN_ORD_RE = re.compile(r"^[一二三四五六七八九十]+、")
PAREN_RE = re.compile(r"^（[0-9]+）|^\([0-9]+\)")

# 重点加粗 token 解析：Markdown 粗体 或 「术语：内容」
BOLD_TOKEN_RE = re.compile(
    r"(?P<md>\*\*(?P<mdtxt>.+?)\*\*)|(?P<colon>(?P<term>[^\s，。；：:！？（）()【】""'<>《》、]{1,14})[：:])"
)

# 表格分隔单元格（如 --- / :-- / :-: / --:）
TABLE_SEP_CELL_RE = re.compile(r"^:?-+:?$")

# 正文引用标注（论文顺序编码制）：[1] / [1,2] / [1-3] / [1, 2-4]
# 仅匹配「[数字…]」形态——公式片段已在 _iter_segments 中被隔离（不会误伤 \sqrt[3]{x}），
# 且 [图1] / [选项A] 这类带文字的方括号标签因非纯数字而不被匹配。
CITE_RE = re.compile(r"\[\s*\d+(?:(?:[,\s、]+|\s*[-–—]\s*)\d+)*\s*\]")

# 参考文献标题识别（用于自动整理参考文献列表）。
# 注意：仅匹配明确的参考文献标题，避免误伤「文献综述」等。
REF_KW = ("参考文献", "引用文献", "参考资料", "文献参考",
          "references", "bibliography", "reference list")


def _is_ref_heading(text: str) -> bool:
    """判断一段标题文本是否「参考文献」类标题。"""
    s = text.strip().rstrip("：: ").strip()
    low = s.lower()
    return any(low == kw or low.startswith(kw) for kw in REF_KW)


# 去掉参考文献条目开头手写的编号（[1] / 1. / 1、 / 1) / (1) / ① 等），
# 交给定制化的自动编号逻辑，避免重复或样式不一致。
_STRIP_REF_NUM_RE = re.compile(
    r"^\s*(?:\[\d+\]\s*"                         # [1]
    r"|\d+[\.、)）]\s*"                          # 1. / 1、 / 1) / 1）
    r"|\(\d+\)\s*"                               # (1)
    r"|[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮]\s*)"               # 带圈数字
)


def _strip_ref_number(s: str) -> str:
    """去掉条目开头的手写编号，保留作者/题名等内容。"""
    return _STRIP_REF_NUM_RE.sub("", s).strip()


# ---------------------------------------------------------------------------
# 数学公式（LaTeX 风格）
# 块级：$$...$$ / \[...\]；行内：$...$ / \(...\)
# 这些片段在解析与排版时一律当成「不透明内容」，避免被加粗 / 标题 / 段落合并规则破坏。
# ---------------------------------------------------------------------------

# 匹配各类公式定界符。行内 $...$ 需经 _looks_like_math 二次判定，避免与货币 $ 混淆。
MATH_RE = re.compile(
    r"\$\$(?P<block>.*?)\$\$"              # 块级 $$
    r"|\\[(?P<brack>.*?)\\]"              # 块级 \[...\]
    r"|\\\((?P<paren>.*?)\\\)"            # 行内 \(...\)
    r"|\$(?P<inline>(?:\\.|[^$\n])+?)\$",  # 行内 $...$
    re.DOTALL,
)


def _looks_like_math(s: str) -> bool:
    """判断一段行内 $...$ 内容是否像数学（避免把 $5 / $100 这类金额当公式）。"""
    s = s.strip()
    if re.search(r"[\\^_{}]", s):                        # 含 LaTeX 命令/上下标/分组
        return True
    if s.isalpha() and len(s) <= 3:                      # 单个短变量 $x$ / $ab$
        return True
    if re.search(r"[=<>+\-*/]", s) and re.search(r"[A-Za-z]", s):  # 含运算符且有字母
        return True
    return False


def _iter_segments(text: str, math_mode: bool = True):
    """产出 (片段文本, 是否公式)。公式片段在后续加粗 / 间距处理中保持不透明。"""
    if not math_mode:
        yield (text, False)
        return
    last = 0
    for m in MATH_RE.finditer(text):
        if m.start() > last:
            yield (text[last:m.start()], False)
        g = m.lastgroup
        if g in ("block", "brack", "paren"):
            is_math = True
        else:  # inline
            is_math = _looks_like_math(m.group("inline") or "")
        yield (m.group(0), is_math)
        last = m.end()
    if last < len(text):
        yield (text[last:], False)


def iter_runs(text: str, math_mode: bool = True, cite_mode: bool = True):
    """产出 (片段, 是否加粗, 是否公式, 是否引用标注)，供 add_rich / 表格单元格使用。"""
    segs = []
    for content, is_math in _iter_segments(text, math_mode):
        if is_math:
            segs.append((content, False, True, False))
        else:
            for seg, bold in tokenize_bold(content):
                if cite_mode:
                    segs.extend(_split_cite(seg, bold))
                else:
                    segs.append((seg, bold, False, False))
    return segs


def _split_cite(seg: str, bold: bool):
    """在普通片段内把「[1]」类引用标注切分为独立 run（保留原文，仅标记）。"""
    out = []
    last = 0
    for m in CITE_RE.finditer(seg):
        if m.start() > last:
            out.append((seg[last:m.start()], bold, False, False))
        out.append((m.group(0), bold, False, True))
        last = m.end()
    if last < len(seg):
        out.append((seg[last:], bold, False, False))
    return out


def _strip_math_delim(span: str) -> str:
    """去掉公式定界符（美化时），保留内部 LaTeX 源码。"""
    s = span
    if len(s) >= 4 and s.startswith("$$") and s.endswith("$$"):
        return s[2:-2]
    if len(s) >= 4 and s.startswith("\\[") and s.endswith("\\]"):
        return s[2:-2]
    if len(s) >= 4 and s.startswith("\\(") and s.endswith("\\)"):
        return s[2:-2]
    if len(s) >= 2 and s.startswith("$") and s.endswith("$"):
        return s[1:-1]
    return s


def _line_is_pure_math(s: str) -> bool:
    """整行是否仅由一个（行内风格）公式组成（不含其它普通文本）。"""
    m = MATH_RE.fullmatch(s)
    if not m:
        return False
    g = m.lastgroup
    if g in ("block", "brack", "paren"):
        return True
    return _looks_like_math(m.group("inline") or "")


def _parse_math_block(lines: list[str], i: int, n: int):
    r"""解析从 i 开始的块级公式（$$...$$ 或 \[...\]，可多行），返回 (block, 下一索引)。"""
    first = lines[i]
    if first.lstrip().startswith("\\["):
        open_d, close_d = "\\[", "\\]"
    else:
        open_d, close_d = "$$", "$$"
    text = first
    open_count = first.count(open_d)
    close_count = first.count(close_d)
    if open_count >= 1 and close_count >= 1 and first.find(open_d) < first.rfind(close_d):
        j = i  # 单行块级公式
    else:
        j = i + 1
        while j < n:
            text += "\n" + lines[j]
            if close_d in lines[j]:
                break
            j += 1
    open_idx = text.find(open_d)
    close_idx = text.rfind(close_d)
    if open_idx == -1 or close_idx == -1 or close_idx <= open_idx:
        inner = ""
    else:
        inner = text[open_idx + len(open_d): close_idx]
    inner = inner.strip("\n")
    block = {"type": "math", "display": True, "text": inner}
    return block, j + 1


# ---------------------------------------------------------------------------
# 文本规范化
# ---------------------------------------------------------------------------

def normalize_text(text: str) -> str:
    """统一换行符、全角空格、去除多余控制字符。"""
    if text is None:
        return ""
    # 统一为 \n
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # 全角空格 / 不间断空格 -> 普通空格
    text = text.replace("\u3000", " ").replace("\xa0", " ")
    # 去掉行内首尾的不可见控制字符（保留 \n）
    lines = []
    for ln in text.split("\n"):
        ln = ln.strip("\x00\x08\x0b\x0c\x1c\x1d\x1e\x1f")
        lines.append(ln.rstrip())
    text = "\n".join(lines)
    # 合并超过 2 个的连续空行为 1 个（最终由解析器再处理）
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


# ---------------------------------------------------------------------------
# 表格（Markdown GFM 表格：表头行 + 分隔行 + 对齐）
# ---------------------------------------------------------------------------

def _split_table_row(line: str) -> list[str]:
    """把一个表格行按 | 切分为单元格（去除首尾空管道）。"""
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _is_table_sep(line: str) -> bool:
    """判断一行是否为 Markdown 表格分隔行（如 |---|:--:|--:|）。"""
    s = line.strip()
    if "|" not in s:
        return False
    cells = _split_table_row(line)
    if not any("-" in c for c in cells):
        return False
    return all(TABLE_SEP_CELL_RE.match(c) for c in cells)


def _table_aligns(sep_line: str) -> list:
    """从分隔行推导每列的居左/居中/居右对齐。"""
    aligns = []
    for c in _split_table_row(sep_line):
        left = c.startswith(":")
        right = c.endswith(":")
        if left and right:
            aligns.append("center")
        elif right:
            aligns.append("right")
        elif left:
            aligns.append("left")
        else:
            aligns.append(None)
    return aligns


def _is_table_start(lines: list[str], i: int, n: int) -> bool:
    """判断从 i 开始是否为一个 Markdown 表格（i 表头行，i+1 分隔行）。"""
    if i + 1 >= n:
        return False
    if "|" not in lines[i]:
        return False
    if not _is_table_sep(lines[i + 1]):
        return False
    return len(_split_table_row(lines[i])) >= 1


def _parse_table(lines: list[str], i: int, n: int):
    """解析从 i 开始的 Markdown 表格，返回 (block, 下一个 i)。"""
    header = _split_table_row(lines[i])
    aligns = _table_aligns(lines[i + 1])
    j = i + 2
    rows = []
    while j < n:
        s = lines[j].strip()
        if s == "":
            break
        if "|" not in s:
            break
        rows.append(_split_table_row(lines[j]))
        j += 1
    block = {"type": "table", "header": header, "aligns": aligns, "rows": rows}
    return block, j


# ---------------------------------------------------------------------------
# 结构解析
# ---------------------------------------------------------------------------

def _looks_like_heading(line: str, smart: bool, next_line: str) -> tuple[bool, int]:
    """返回 (是否标题, 级别)。普通文本的智能识别仅在 smart=True 时启用。"""
    if not line.strip():
        return (False, 0)

    # 1) Markdown 标题永远尊重
    m = MD_HEADING_RE.match(line)
    if m:
        lvl = min(len(m.group(1)), 3)  # 超过 3 级统一按 3 级处理
        return (True, lvl)

    # 2) 明确的章节 / 编号模式
    if CHAPTER_RE.match(line) or DOTNUM_RE.match(line) or CN_ORD_RE.match(line) or PAREN_RE.match(line):
        # 章节/编号 -> 1 级；子编号(含点) -> 2 级
        if DOTNUM_RE.match(line) and "." in line[:6]:
            return (True, 2)
        return (True, 1)

    # 3) 智能识别：短行、无句末标点，且下一行明显更长（像正文）
    if smart:
        stripped = line.strip()
        if (len(stripped) <= 18
                and not stripped[-1] in WEAK_END
                and not BULLET_RE.match(line)
                and not ORDERED_RE.match(line)):
            nxt = next_line.strip()
            if nxt and len(nxt) >= 2 * len(stripped) and len(nxt) >= 12:
                return (True, 2)
    return (False, 0)


def _is_title_like(s: str) -> bool:
    """首行是否像文档主标题：较短、无句末标点、非列表/标题标记。"""
    if not s or len(s) > 20:
        return False
    if s[-1] in WEAK_END:
        return False
    if BULLET_RE.match(s) or ORDERED_RE.match(s) or MD_HEADING_RE.match(s):
        return False
    return True


def parse_document(text: str, smart_heading: bool = True) -> list[dict]:
    """
    将文本解析为结构化 block 列表。
    每个 block:
      {'type':'title',     'text':...}
      {'type':'heading',   'level':1..3, 'text':...}
      {'type':'paragraph', 'text':...}
      {'type':'list',      'ordered':bool, 'items':[...]}
      {'type':'table',     'header':[...], 'aligns':[...], 'rows':[[...], ...]}
      {'type':'math',      'display':True, 'text':...}   # 块级公式（LaTeX 源码）
      {'type':'reference', 'entries':[...]}              # 参考文献条目（位于「参考文献」标题下）
      {'type':'hr'}
    """
    text = normalize_text(text)
    raw_lines = text.split("\n")

    # 先清洗空行，得到非空行序列（保留原始用于上下文）
    lines = [ln for ln in raw_lines]

    blocks: list[dict] = []
    i = 0
    n = len(lines)
    title_set = False
    in_refs = False  # 是否已进入「参考文献」区域（遇到非参考文献标题即结束）

    # 缓冲
    para_buf: list[str] = []
    list_buf: list[str] = []
    list_ordered = False

    def flush_para():
        if para_buf:
            joined = _merge_paragraph(para_buf)
            if joined.strip():
                blocks.append({"type": "paragraph", "text": joined.strip()})
            para_buf.clear()

    def flush_list():
        if list_buf:
            blocks.append({"type": "list", "ordered": list_ordered, "items": list(list_buf)})
            list_buf.clear()

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 空行：段落 / 列表的分隔
        if stripped == "":
            flush_para()
            flush_list()
            i += 1
            continue

        # 水平分隔线
        if re.fullmatch(r"[-—=*_]{3,}", stripped):
            flush_para(); flush_list()
            blocks.append({"type": "hr"})
            i += 1
            continue

        # 块级数学公式 $$...$$ 或 \[...\]（多行，居中独立段落，原样保留）
        if stripped.startswith("$$") or stripped.startswith("\\["):
            flush_para(); flush_list()
            block, i = _parse_math_block(lines, i, n)
            blocks.append(block)
            continue

        # 表格（Markdown GFM 表格：表头行 + 分隔行）
        if _is_table_start(lines, i, n):
            flush_para(); flush_list()
            block, i = _parse_table(lines, i, n)
            blocks.append(block)
            continue

        # 列表项（无序）
        bm = BULLET_RE.match(line)
        if bm:
            flush_para()
            if in_refs:
                blocks.append({"type": "reference", "entries": [bm.group(2).strip()]})
                i += 1
                continue
            list_ordered = False
            list_buf.append(bm.group(2).strip())
            i += 1
            continue

        # 有序项：仅当连续出现多个时才作为有序列表，否则当作编号标题
        om = ORDERED_RE.match(line)
        if om:
            items = [om.group(2).strip()]
            j = i + 1
            while j < n:
                lj = lines[j].strip()
                if lj == "":
                    j += 1
                    continue
                omj = ORDERED_RE.match(lines[j])
                if omj:
                    items.append(omj.group(2).strip())
                    j += 1
                    continue
                break
            if len(items) >= 2:
                flush_para(); flush_list()
                if in_refs:
                    blocks.append({"type": "reference", "entries": list(items)})
                else:
                    blocks.append({"type": "list", "ordered": True, "items": items})
                i = j
                continue
            # 单个有序行 -> 落到下方标题识别（作为编号标题）

        # 整行纯公式（行内风格但独立成行）：避免被误判为标题
        if _line_is_pure_math(stripped):
            flush_list()
            para_buf.append(stripped)
            i += 1
            continue

        # 标题
        is_head, lvl = _looks_like_heading(line, smart_heading, lines[i + 1] if i + 1 < n else "")
        if is_head:
            flush_para(); flush_list()
            heading_text = stripped
            # 去掉编号前缀里的标点，保持整洁（保留数字）
            if not MD_HEADING_RE.match(line):
                heading_text = re.sub(r"^[\s（(]*[0-9]+[\.、）)]*\s*", "", heading_text)
                heading_text = re.sub(r"^[一二三四五六七八九十]+、", "", heading_text)
            is_chapter = bool(CHAPTER_RE.match(line) or DOTNUM_RE.match(line)
                              or CN_ORD_RE.match(line) or PAREN_RE.match(line))
            # 参考文献标题：开启引用区标记；其它标题：关闭（引用区结束）
            in_refs = _is_ref_heading(heading_text)
            if not title_set and lvl == 1 and len(blocks) == 0 and not is_chapter:
                # 文档首个非编号标题 -> 作为文档主标题
                blocks.append({"type": "title", "text": heading_text})
                title_set = True
            else:
                blocks.append({"type": "heading", "level": lvl, "text": heading_text})
            i += 1
            continue

        # 普通正文行（首行且尚无标题时，提升为文档主标题）
        if not title_set and len(blocks) == 0 and _is_title_like(stripped):
            flush_list()
            blocks.append({"type": "title", "text": stripped})
            title_set = True
            i += 1
            continue
        if in_refs:
            # 参考文献区域：每行独立成条，交给 build_docx 统一悬挂缩进+编号
            flush_para(); flush_list()
            blocks.append({"type": "reference", "entries": [stripped]})
            i += 1
            continue
        flush_list()
        para_buf.append(stripped)
        i += 1

    flush_para(); flush_list()
    return blocks


def _merge_plain(text: str) -> str:
    """对不含公式的纯文本片段做换行合并与空格清理。"""
    lines = text.split("\n")
    out_parts = []
    for idx, ln in enumerate(lines):
        ln = ln.strip()
        if idx == 0:
            out_parts.append(ln)
            continue
        prev = out_parts[-1]
        # 上一行末尾是拉丁字母/数字，且本行开头也是 -> 补一个空格
        if prev and prev[-1].isascii() and prev[-1].isalnum() and ln and ln[0].isascii() and ln[0].isalnum():
            out_parts.append(" " + ln)
        else:
            out_parts.append(ln)
    text = "".join(out_parts)
    # 清理多余空格（中文间不留空格）
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\s+([，。；：！？、）】》])", r"\1", text)
    text = re.sub(r"([（【《])\s+", r"\1", text)
    return text


def _merge_paragraph(lines: list[str]) -> str:
    """将一段内被换行拆散的句子重新合并；公式片段保持不透明，不被插入空格。"""
    raw = "\n".join(ln.strip() for ln in lines)
    out = []
    for content, is_math in _iter_segments(raw, math_mode=True):
        if is_math:
            out.append(content)
        else:
            out.append(_merge_plain(content))
    return "".join(out)


# ---------------------------------------------------------------------------
# 重点加粗 token 化
# ---------------------------------------------------------------------------

def tokenize_bold(text: str) -> list[tuple[str, bool]]:
    """
    将一段文本切分为 (片段, 是否加粗) 的 token 列表。
    规则：Markdown `**x**` 加粗；「短术语：内容」中冒号前的术语加粗。
    """
    tokens: list[tuple[str, bool]] = []
    pos = 0
    for m in BOLD_TOKEN_RE.finditer(text):
        if m.start() > pos:
            tokens.append((text[pos:m.start()], False))
        if m.group("md"):
            tokens.append((m.group("mdtxt"), True))
        else:  # colon
            term = m.group("term")
            tokens.append((term, True))
            tokens.append((m.group(0)[len(term):], False))  # 冒号本身
        pos = m.end()
    if pos < len(text):
        tokens.append((text[pos:], False))
    return tokens


# ---------------------------------------------------------------------------
# 结构大纲（用于预览）
# ---------------------------------------------------------------------------

def outline(blocks: list[dict]) -> list[tuple[int, str]]:
    """返回 (级别, 文本) 列表，仅包含 title/heading。"""
    res = []
    for b in blocks:
        if b["type"] == "title":
            res.append((0, b["text"]))
        elif b["type"] == "heading":
            res.append((b["level"], b["text"]))
    return res


# ---------------------------------------------------------------------------
# 样式预设
# ---------------------------------------------------------------------------

# 每个预设：body 字体/字号、heading 字体、行距倍数、段后间距(pt)、页边距(cm)
STYLE_PRESETS = {
    "默认": {
        "body_font": "宋体", "body_size": 11,
        "head_font": "微软雅黑", "title_font": "微软雅黑",
        "latin_font": "Times New Roman",
        "line_spacing": 1.4, "space_after": 6, "margin_cm": 2.5,
        "title_center": True,
    },
    "毕业论文": {
        # 严谨学位论文：正文宋体小四、标题黑体、行距 1.5、页边距 2.5cm、标题居中
        "body_font": "宋体", "body_size": 12,
        "head_font": "黑体", "title_font": "黑体",
        "latin_font": "Times New Roman",
        "line_spacing": 1.5, "space_after": 8, "margin_cm": 2.5,
        "title_center": True,
    },
    "通用论文": {
        # 一般学术论文：正文宋体小四、标题黑体、行距 1.4、页边距 2.5cm
        "body_font": "宋体", "body_size": 11,
        "head_font": "黑体", "title_font": "黑体",
        "latin_font": "Times New Roman",
        "line_spacing": 1.4, "space_after": 6, "margin_cm": 2.5,
        "title_center": True,
    },
    "课程论文": {
        # 课程作业：略宽松、标题用微软雅黑、页边距 2.3cm
        "body_font": "宋体", "body_size": 11,
        "head_font": "微软雅黑", "title_font": "微软雅黑",
        "latin_font": "Times New Roman",
        "line_spacing": 1.35, "space_after": 6, "margin_cm": 2.3,
        "title_center": True,
    },
    "简洁风": {
        "body_font": "等线", "body_size": 11,
        "head_font": "微软雅黑", "title_font": "微软雅黑",
        "latin_font": "Times New Roman",
        "line_spacing": 1.25, "space_after": 4, "margin_cm": 2.2,
        "title_center": False,
    },
    "阅读风": {
        "body_font": "宋体", "body_size": 13,
        "head_font": "微软雅黑", "title_font": "微软雅黑",
        "latin_font": "Times New Roman",
        "line_spacing": 1.6, "space_after": 10, "margin_cm": 3.0,
        "title_center": True,
    },
}


# ---------------------------------------------------------------------------
# 生成 Word
# ---------------------------------------------------------------------------

def build_docx(blocks: list[dict], out_path: str, preset_name: str = "默认",
               auto_toc: bool = False, math_pretty: bool = True,
               cite_sup: bool = True, header_text: str = None,
               footer_mode: str = "none", footer_text: str = "",
               footer_align: str = "center",
               ref_auto: bool = True,
               ref_style: str = "gb7714", ref_hang: float = 0.74,
               ref_line: float = 1.5,
               body_font: str = None, head_font: str = None,
               title_font: str = None, body_size: float = None,
               latin_font: str = None) -> None:
    """把结构化 block 列表写成 .docx 文件。

    header_text: 页眉文字（留空/None 则不添加页眉）。
    footer_mode: 页脚内容类型（详见 _set_footer：无 / 页码 /
                 第X页共Y页 / 日期 / 日期+页码 / 自定义文字 / 左文字右页码）。
    footer_text: 页脚自定义文字（供 text / text_page_split 模式使用）。
    footer_align: 页脚对齐方式（left / center / right），与 footer_mode 解耦，
                 仅对单元素页脚（页码/第X页共Y页/日期/自定义文字）生效。
    ref_auto: 是否自动整理「参考文献」标题下的条目（悬挂缩进+自动编号）。
    ref_style: 参考文献编号样式，详见 _add_references（gb7714 / numbered /
                 paren / circle / superscript / none）。
    ref_hang: 悬挂缩进量（厘米）。
    ref_line: 参考文献行距倍数。
    body_font / head_font / title_font / body_size: 覆盖预设的字体与正文字号；
                 为 None 时回退到预设默认值。
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    _base = STYLE_PRESETS.get(preset_name, STYLE_PRESETS["默认"])
    # 复制预设并以界面显式选择的字体/字号覆盖，保证「预设=快速套用、可被单独调整」
    preset = dict(_base)
    if body_font:
        preset["body_font"] = body_font      # 仅覆盖中文字体（eastAsia）
    if head_font:
        preset["head_font"] = head_font
    if title_font:
        preset["title_font"] = title_font
    if body_size is not None:
        preset["body_size"] = body_size
    if latin_font:
        preset["latin_font"] = latin_font    # 西文字体（ascii/hAnsi）单独控制，默认 Times New Roman

    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(preset["margin_cm"])
        section.bottom_margin = Cm(preset["margin_cm"])
        section.left_margin = Cm(preset["margin_cm"])
        section.right_margin = Cm(preset["margin_cm"])

    # 页眉 / 页脚
    if header_text and header_text.strip():
        _set_header(doc, header_text.strip())
    if footer_mode and footer_mode != "none":
        _set_footer(doc, footer_mode, footer_text or "", footer_align or "center")

    # 默认正文字体（中英混排：西文→latin_font，中文→body_font）
    normal = doc.styles["Normal"]
    normal.font.name = preset["latin_font"]
    normal.font.size = Pt(preset["body_size"])
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), preset["body_font"])
    normal.paragraph_format.line_spacing = preset["line_spacing"]
    normal.paragraph_format.space_after = Pt(preset["space_after"])

    # 标题样式字体
    for lvl in (0, 1, 2, 3):
        try:
            st = doc.styles[f"Heading {lvl}"] if lvl > 0 else doc.styles["Title"]
        except KeyError:
            continue
        fname = preset["title_font"] if lvl == 0 else preset["head_font"]
        st.font.name = preset["latin_font"]
        st.font.size = Pt(preset["body_size"] + (6 if lvl == 0 else (4 - lvl)))
        st.element.rPr.rFonts.set(qn("w:eastAsia"), fname)

    def add_rich(paragraph, text):
        for seg, bold, is_math, is_cite in iter_runs(text, math_mode=True, cite_mode=cite_sup):
            if not seg:
                continue
            if is_math:
                disp = _strip_math_delim(seg) if math_pretty else seg
                run = paragraph.add_run(disp)
                if math_pretty:
                    run.italic = True
                run.font.name = preset["latin_font"]
                run._element.rPr.rFonts.set(qn("w:eastAsia"), preset["body_font"])
            else:
                run = paragraph.add_run(seg)
                run.bold = bold
                # 确保 rPr 存在，并设置中文字体
                run.font.name = preset["latin_font"]
                run._element.rPr.rFonts.set(qn("w:eastAsia"), preset["body_font"])
            if is_cite:
                run.font.superscript = True

    # 自动目录
    if auto_toc:
        _insert_toc(doc)

    ref_counter = [0]  # 参考文献连续编号（跨多个 reference block 累加）

    for b in blocks:
        t = b["type"]
        if t == "title":
            p = doc.add_heading(b["text"], level=0)
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                           if preset["title_center"] else WD_ALIGN_PARAGRAPH.LEFT)
        elif t == "heading":
            p = doc.add_heading(b["text"], level=b["level"])
        elif t == "hr":
            p = doc.add_paragraph()
            p.add_run("—" * 18).font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        elif t == "list":
            style = "List Number" if b.get("ordered") else "List Bullet"
            for item in b["items"]:
                p = doc.add_paragraph(style=style)
                add_rich(p, item)
        elif t == "math":
            _add_math_block(doc, b, preset, math_pretty)
        elif t == "table":
            _add_table(doc, b, preset, math_pretty, cite_sup)
        elif t == "reference":
            _add_references(doc, b, preset, ref_auto, ref_style, ref_hang,
                           ref_line, math_pretty, cite_sup, add_rich, ref_counter)
        elif t == "paragraph":
            p = doc.add_paragraph()
            add_rich(p, b["text"])

    # 导出前强制全黑：标题/分隔线/目录条目/表格边框一律纯黑，不保留任何主题色
    _force_all_black(doc)
    doc.save(out_path)


def _force_all_black(doc):
    """把整篇文档（正文 + 全部样式 part + 页眉页脚）的所有颜色强制为纯黑。

    python-docx 默认模板里 Heading/Title/Subtitle/Caption 等样式带
    蓝色系主题色（accent1/text2），只改字体字号时标题会导出成蓝色。
    这里遍历包内所有 XML part 中的 w:color 元素：val 一律改 000000，
    并删掉 themeColor/themeShade/themeTint 属性（主题色引用优先级高于
    val，不删的话 Word 仍会按主题色渲染）。

    注意除了 word/styles.xml，模板还带 word/stylesWithEffects.xml
    （Word 实际读取的带特效样式副本），所以必须遍历包内全部 part，
    不能只改 doc.styles。
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _fix(root):
        if root is None:
            return
        for el in root.iter():
            if el.tag == qn("w:color"):
                el.set(qn("w:val"), "000000")
                for attr in ("w:themeColor", "w:themeShade", "w:themeTint"):
                    if el.get(qn(attr)) is not None:
                        del el.attrib[qn(attr)]

    # 遍历包内所有 XML part：styles / document / 页眉页脚 / numbering 等；
    # 二进制 part（图片）没有 _element，自动跳过
    for part in doc.part.package.parts:
        root = getattr(part, "_element", None)
        if root is None:
            root = getattr(part, "element", None)
        try:
            _fix(root)
        except Exception:
            pass

    # word/stylesWithEffects.xml 被 python-docx 当普通二进制 Part 加载
    # （无 _element），但 Word 打开文档时会优先读它——必须单独修字节流。
    import re as _re
    for part in doc.part.package.parts:
        if "stylesWithEffects" not in str(part.partname):
            continue
        blob = getattr(part, "_blob", None)
        if not blob:
            continue
        xml = blob.decode("utf-8")

        def _sub_color(m):
            tag = m.group(0)
            tag = _re.sub(r'w:val="[^"]*"', 'w:val="000000"', tag)
            tag = _re.sub(r'\s+w:theme(?:Color|Shade|Tint)="[^"]*"', "", tag)
            return tag

        xml = _re.sub(r"<w:color\b[^>]*/?>", _sub_color, xml)
        part._blob = xml.encode("utf-8")

    # Word 更新 TOC 域时条目套用 Hyperlink 字符样式（内置默认蓝色）。
    # 模板里没有该样式就补一个黑色定义，保证目录条目也是黑色。
    styles_el = doc.styles.element
    has_hyperlink = any(
        st.get(qn("w:styleId")) == "Hyperlink"
        for st in styles_el.findall(qn("w:style"))
    )
    if not has_hyperlink:
        st = OxmlElement("w:style")
        st.set(qn("w:type"), "character")
        st.set(qn("w:styleId"), "Hyperlink")
        name = OxmlElement("w:name")
        name.set(qn("w:val"), "Hyperlink")
        st.append(name)
        rPr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "000000")
        rPr.append(color)
        st.append(rPr)
        styles_el.append(st)


def _align_para(p, align):
    """设置段落水平对齐：left / center / right / None(默认左)。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _set_table_borders(table):
    """给表格所有边框加单线（不依赖内置样式，确保可见边框）。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def _add_table(doc, b, preset, math_pretty: bool = True, cite_sup: bool = True):
    """把 table block 渲染为带边框 Word 表格：表头加粗 + 列对齐 + 预设字体；单元格支持公式。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    header = list(b.get("header", []))
    aligns = list(b.get("aligns", []))
    rows = [list(r) for r in b.get("rows", [])]

    n_cols = max(len(header), (max((len(r) for r in rows), default=0)), 1)

    def norm(cells):
        cells = list(cells)
        while len(cells) < n_cols:
            cells.append("")
        return cells[:n_cols]

    header = norm(header)
    rows = [norm(r) for r in rows]

    table = doc.add_table(rows=1, cols=n_cols)
    _set_table_borders(table)

    body_size = preset["body_size"]

    # 表头行
    for idx, cell in enumerate(table.rows[0].cells):
        cell.text = header[idx]
        p = cell.paragraphs[0]
        _align_para(p, aligns[idx] if idx < len(aligns) else None)
        for run in p.runs:
            run.bold = True
            run.font.name = preset["latin_font"]
            run.font.size = Pt(body_size)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), preset["head_font"])

    # 数据行（单元格支持内联加粗）
    for r in rows:
        cells = table.add_row().cells
        for idx, cell in enumerate(cells):
            p = cell.paragraphs[0]
            # 清空表格默认空 run，避免残留空片段
            for rr in list(p.runs):
                rr._element.getparent().remove(rr._element)
            _align_para(p, aligns[idx] if idx < len(aligns) else None)
            for seg, bold, is_math, is_cite in iter_runs(r[idx], math_mode=True, cite_mode=cite_sup):
                if not seg:
                    continue
                if is_math:
                    disp = _strip_math_delim(seg) if math_pretty else seg
                    run = p.add_run(disp)
                    if math_pretty:
                        run.italic = True
                else:
                    run = p.add_run(seg)
                    run.bold = bold
                run.font.name = preset["latin_font"]
                run.font.size = Pt(body_size)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), preset["body_font"])
                if is_cite:
                    run.font.superscript = True


def _add_math_block(doc, b, preset, math_pretty: bool = True):
    """渲染块级公式：居中独立段落，保留换行；美化时整体斜体。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    content = b.get("text", "") or ""
    p = doc.add_paragraph()
    _align_para(p, "center")
    # 公式段落不加段后间距，避免与正文过分开裂
    p.paragraph_format.space_after = Pt(0)
    lines = content.split("\n")
    for k, line in enumerate(lines):
        run = p.add_run(line)
        if math_pretty:
            run.italic = True
        run.font.name = preset["latin_font"]
        run.font.size = Pt(preset["body_size"])
        run._element.rPr.rFonts.set(qn("w:eastAsia"), preset["body_font"])
        if k < len(lines) - 1:
            run.add_break()


def _add_references(doc, b, preset, ref_auto: bool, ref_style: str,
                    ref_hang: float, ref_line: float, math_pretty: bool,
                    cite_sup: bool, add_rich, counter: list):
    """渲染参考文献条目。

    - ref_auto 关闭：原样逐条输出（保留用户手写的编号/格式）。
    - ref_auto 开启：统一悬挂缩进 + 自动连续编号（编号样式可切换），
      并剥离条目开头的手写编号，避免重复。
    """
    from docx.shared import Cm, Pt

    entries = b.get("entries", [])
    for raw in entries:
        counter[0] += 1
        if not ref_auto:
            p = doc.add_paragraph()
            add_rich(p, raw)
            continue
        txt = _strip_ref_number(raw)
        p = doc.add_paragraph()
        # 悬挂缩进：首行回退到左边界外，后续行缩进 ref_hang
        p.paragraph_format.left_indent = Cm(ref_hang)
        p.paragraph_format.first_line_indent = Cm(-ref_hang)
        p.paragraph_format.line_spacing = ref_line
        p.paragraph_format.space_after = Pt(preset["space_after"])
        # 根据编号样式添加前缀标记
        n = counter[0]
        if ref_style == "gb7714":
            p.add_run(f"[{n}] ")
        elif ref_style == "paren":
            p.add_run(f"({n}) ")
        elif ref_style == "circle":
            if 1 <= n <= len(CIRCLED):
                p.add_run(CIRCLED[n - 1] + " ")
            else:
                p.add_run(f"[{n}] ")
        elif ref_style == "superscript":
            r = p.add_run(f"{n}")
            r.font.superscript = True
            p.add_run("  ")
        elif ref_style == "none":
            pass  # 仅悬挂缩进，不加编号（适合著者-出版年制，由用户自行书写）
        else:  # numbered（默认）
            p.add_run(f"{n}. ")
        add_rich(p, txt)


def _insert_field(paragraph, instr: str):
    """在段落末尾插入一个 Word 域（如 PAGE / NUMPAGES），用于自动页码。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = instr
    fldSep = OxmlElement("w:fldChar")
    fldSep.set(qn("w:fldCharType"), "separate")
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    r = run._element
    r.append(fldBegin)
    r.append(instrText)
    r.append(fldSep)
    r.append(fldEnd)
    return run


def _clear_para_runs(paragraph):
    """删除段落里所有 run（用于清空页眉/页脚默认空段落）。"""
    for rr in list(paragraph.runs):
        rr._element.getparent().remove(rr._element)


def _set_header(doc, text: str):
    """为文档所有节添加（或覆盖）页眉文字，左对齐。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for section in doc.sections:
        p = section.header.paragraphs[0]
        _clear_para_runs(p)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(text)


def _set_footer(doc, mode: str, footer_text: str = "", align: str = "center"):
    """为文档所有节添加页脚。页脚「内容类型」与「对齐方式」解耦：
    mode 决定显示什么，align（left/center/right）单独决定横向位置。

    mode 取值：
      none           无页脚（调用方已拦截，此处兜底）
      page           页码（位置由 align 决定）
      page_of_total  第 X 页 / 共 Y 页（位置由 align 决定）
      date           日期（位置由 align 决定）
      date_page      左日期 / 右页码（固定分栏，align 不生效）
      text           自定义文字（位置由 align 决定，内容取 footer_text）
      text_page_split 左自定义文字 / 右页码（固定分栏，align 不生效）
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.shared import Cm

    _ALIGN = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }

    today = datetime.date.today().strftime("%Y-%m-%d")

    def _usable_cm(section):
        # 页脚可用宽度（页面宽 - 左右页边距），单位 cm，用于右对齐制表位
        return max(1.0, (section.page_width - section.left_margin
                         - section.right_margin) / 360000.0)

    def _add_page(p):
        _insert_field(p, "PAGE")

    def _add_total(p):
        _insert_field(p, "NUMPAGES")

    def _add_date(p):
        p.add_run(today)

    def _split(p, section, left_fn, right_fn):
        # 左内容 + 制表位到行尾右对齐的右内容（日期+页码 / 文字+页码）
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(_usable_cm(section)), WD_TAB_ALIGNMENT.RIGHT)
        left_fn(p)
        p.add_run("\t")
        right_fn(p)

    for section in doc.sections:
        p = section.footer.paragraphs[0]
        _clear_para_runs(p)
        if mode == "page":
            p.alignment = _ALIGN.get(align, WD_ALIGN_PARAGRAPH.CENTER)
            _add_page(p)
        elif mode == "page_of_total":
            p.alignment = _ALIGN.get(align, WD_ALIGN_PARAGRAPH.CENTER)
            p.add_run("第 "); _add_page(p); p.add_run(" 页 / 共 ")
            _add_total(p); p.add_run(" 页")
        elif mode == "date":
            p.alignment = _ALIGN.get(align, WD_ALIGN_PARAGRAPH.CENTER)
            _add_date(p)
        elif mode == "date_page":
            _split(p, section, _add_date, _add_page)
        elif mode == "text":
            p.alignment = _ALIGN.get(align, WD_ALIGN_PARAGRAPH.CENTER)
            p.add_run(footer_text or "")
        elif mode == "text_page_split":
            _split(p, section, lambda q: q.add_run(footer_text or ""), _add_page)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_page(p)


def _insert_toc(doc):
    """插入一个 Word 域代码式目录（打开文档时 Word 会提示更新）。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldSep = OxmlElement("w:fldChar")
    fldSep.set(qn("w:fldCharType"), "separate")
    fldText = OxmlElement("w:t")
    fldText.text = "【右键 -> 更新域，可生成目录】"
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    r = run._element
    r.append(fldBegin)
    r.append(instrText)
    r.append(fldSep)
    r.append(fldText)
    r.append(fldEnd)
    doc.add_paragraph()  # 目录与正文空一行


def format_text_to_docx(text: str, out_path: str, preset_name: str = "默认",
                        smart_heading: bool = True, auto_toc: bool = False,
                        math_pretty: bool = True, cite_sup: bool = True,
                        header_text: str = None, footer_mode: str = "none",
                        footer_text: str = "", footer_align: str = "center",
                        ref_auto: bool = True,
                        ref_style: str = "gb7714", ref_hang: float = 0.74,
                        ref_line: float = 1.5,
                        body_font: str = None, head_font: str = None,
                        title_font: str = None, body_size: float = None,
                        latin_font: str = None) -> list[dict]:
    """一站式：文本 -> 解析 -> docx，返回解析出的 blocks（供预览）。"""
    blocks = parse_document(text, smart_heading=smart_heading)
    build_docx(blocks, out_path, preset_name=preset_name, auto_toc=auto_toc,
               math_pretty=math_pretty, cite_sup=cite_sup,
               header_text=header_text, footer_mode=footer_mode,
               footer_text=footer_text, footer_align=footer_align,
               ref_auto=ref_auto, ref_style=ref_style,
               ref_hang=ref_hang, ref_line=ref_line,
               body_font=body_font, head_font=head_font,
               title_font=title_font, body_size=body_size,
               latin_font=latin_font)
    return blocks


# ---------------------------------------------------------------------------
# 简单自测（直接运行本文件时会跑）
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    sample = """我的学习笔记

一、绪论
这是绪论部分的内容。它描述了研究的背景与意义，内容比较长比较长比较长比较长比较长。

1. 研究背景
背景包括很多方面，例如技术发展与社会需求。

2. 研究意义
意义在于提升效率。

重点：这一部分很重要。
定义：模型是指一组参数。

- 苹果
- 香蕉
- 橘子

第二章 方法
方法是本文的核心。继续写一些内容来测试段落合并是否正确，this is English and should keep space.
"""
    blocks = parse_document(sample, smart_heading=True)
    print("=== blocks ===")
    for b in blocks:
        print(b)
    print("=== outline ===")
    for o in outline(blocks):
        print(o)
    print("=== tokens of '重点：这一部分很重要。' ===")
    print(tokenize_bold("重点：这一部分很重要。"))
